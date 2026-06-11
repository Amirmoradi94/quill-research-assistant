"""Document upload, listing, retrieval, and text extraction.

Files are stored under the configured app data directory.
On upload, plain text is extracted (.pdf via pypdf, .docx via python-docx,
.txt/.md/.tex raw) and stored in the ``text`` column so Quill workflows can
read documents without re-parsing them.
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Optional

import asyncio

from fastapi import APIRouter, BackgroundTasks, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session

from .database import get_db
from . import models
from .runtime import documents_dir

router = APIRouter(prefix="/api/documents", tags=["documents"])

DATA_DIR = documents_dir()

ALLOWED_KINDS = {
    "cv", "research_statement", "sample_paper",
    "cover_letter_tmpl", "transcript", "other",
}
MAX_BYTES = 25 * 1024 * 1024  # 25 MB

_slug_re = re.compile(r"[^a-zA-Z0-9._-]+")


def _slugify(name: str) -> str:
    return _slug_re.sub("-", name).strip("-")[:80] or "file"


def _extract_text(path: Path) -> str:
    """Best-effort text extraction. Returns empty string on any failure.

    For PDFs we also pull URL annotations (the underlying hyperlinks attached
    to "GitHub", "LinkedIn", etc.) and append them so the LLM can see real
    URLs that aren't part of the visible text run.
    """
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            body = "\n\n".join((p.extract_text() or "") for p in reader.pages).strip()
            urls = _extract_pdf_urls(reader)
            if urls:
                body += "\n\n--- Hyperlinks in PDF ---\n" + "\n".join(urls)
            return body
        if suffix == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            return "\n".join(p.text for p in doc.paragraphs).strip()
        if suffix in {".txt", ".md", ".markdown", ".tex", ".bib", ".csv", ".json", ".yaml", ".yml"}:
            return path.read_text(encoding="utf-8", errors="replace").strip()
    except Exception:
        return ""
    return ""


def _extract_pdf_urls(reader) -> list[str]:
    """Pull URI actions out of a PDF's link annotations. Deduplicated, in order."""
    seen: set[str] = set()
    out: list[str] = []
    try:
        for page in reader.pages:
            annots = page.get("/Annots")
            if not annots:
                continue
            for a in annots:
                try:
                    obj = a.get_object() if hasattr(a, "get_object") else a
                    action = obj.get("/A") if hasattr(obj, "get") else None
                    if not action:
                        continue
                    uri = action.get("/URI") if hasattr(action, "get") else None
                    if uri:
                        u = str(uri).strip()
                        if u and u not in seen:
                            seen.add(u)
                            out.append(u)
                except Exception:
                    continue
    except Exception:
        return out
    return out


def _doc_to_dict(d: models.Document) -> dict:
    p = Path(d.file_path) if d.file_path else None
    size = p.stat().st_size if p and p.exists() else 0
    return {
        "id": d.id,
        "kind": d.kind,
        "title": d.title,
        "filename": p.name if p else "",
        "extension": (p.suffix.lstrip(".") if p else ""),
        "size_bytes": size,
        "is_default": d.is_default,
        "version": d.version,
        "has_text": bool(d.text and d.text.strip()),
        "text_chars": len(d.text or ""),
        "created_at": d.created_at.isoformat() if d.created_at else None,
        "updated_at": d.updated_at.isoformat() if d.updated_at else None,
    }


@router.get("")
def list_documents(kind: Optional[str] = None, db: Session = Depends(get_db)):
    q = db.query(models.Document)
    if kind:
        q = q.filter(models.Document.kind == kind)
    rows = q.order_by(models.Document.is_default.desc(), models.Document.created_at.desc()).all()
    return [_doc_to_dict(r) for r in rows]


@router.post("", status_code=201)
async def upload_document(
    background_tasks: BackgroundTasks,
    kind: str = Form(...),
    title: Optional[str] = Form(None),
    is_default: bool = Form(False),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    if kind not in ALLOWED_KINDS:
        raise HTTPException(400, f"Invalid kind. Allowed: {sorted(ALLOWED_KINDS)}")
    contents = await file.read()
    if len(contents) > MAX_BYTES:
        raise HTTPException(413, f"File exceeds {MAX_BYTES // (1024*1024)} MB limit.")
    if len(contents) == 0:
        raise HTTPException(400, "Empty file.")

    kind_dir = DATA_DIR / kind
    kind_dir.mkdir(parents=True, exist_ok=True)

    safe_name = _slugify(file.filename or "upload")
    doc = models.Document(
        kind=kind,
        title=title or (file.filename or safe_name),
        is_default=is_default,
        file_path="",
        text="",
    )
    db.add(doc)
    db.flush()

    file_path = kind_dir / f"{doc.id}_{safe_name}"
    file_path.write_bytes(contents)
    doc.file_path = str(file_path)
    doc.text = _extract_text(file_path)

    if is_default:
        db.query(models.Document).filter(
            models.Document.kind == kind,
            models.Document.id != doc.id,
        ).update({"is_default": False})

    db.commit()
    db.refresh(doc)

    # If this is a CV upload, auto-wire it to the user profile and kick off
    # the LLM extraction so the Profile page populates without user action.
    # Same for transcripts — add them to user.transcript_doc_ids.
    _autowire_to_user_profile(db, doc, is_default, background_tasks)

    return _doc_to_dict(doc)


def _autowire_to_user_profile(
    db: Session, doc: models.Document, is_default: bool, background_tasks: BackgroundTasks
) -> None:
    """When a CV or transcript is uploaded, link it to users(id=1) and trigger
    profile extraction in the background. Runs silently — the user just sees
    their Profile page populate over the next ~30s."""
    user = db.get(models.User, 1)
    if not user:
        return

    should_trigger_extraction = False
    if doc.kind == "cv":
        # Set as primary CV if user has none, or if explicitly marked default.
        if user.cv_doc_id is None or is_default:
            user.cv_doc_id = doc.id
            should_trigger_extraction = bool((doc.text or "").strip())
    elif doc.kind == "transcript":
        ids = list(user.transcript_doc_ids or [])
        if doc.id not in ids:
            ids.append(doc.id)
            user.transcript_doc_ids = ids
            should_trigger_extraction = bool(user.cv_doc_id) and bool((doc.text or "").strip())
    elif doc.kind == "research_statement":
        if user.research_statement_doc_id is None or is_default:
            user.research_statement_doc_id = doc.id
    elif doc.kind == "sample_paper":
        ids = list(user.sample_paper_doc_ids or [])
        if doc.id not in ids:
            ids.append(doc.id)
            user.sample_paper_doc_ids = ids

    db.commit()

    if should_trigger_extraction:
        background_tasks.add_task(_run_profile_extraction_silently)


def _run_profile_extraction_silently() -> None:
    """Spawn the extract_user_profile_full workflow and drain its SSE stream
    so the on-success handler in quill.py writes the extracted fields back to
    the DB. We swallow exceptions — failures are visible via /api/ai/runs."""
    from .database import SessionLocal
    from .quill import (
        _settings, _select_provider_for, _resolve_user_context,
        _run_and_stream, RunRequest, Workflow,
    )

    db = SessionLocal()
    try:
        settings = _settings(db)
        provider = _select_provider_for(settings)
        if provider is None:
            return  # no provider configured — silently no-op
        cli_path = (
            settings.claude_cli_path if provider.value == "claude_cli" else
            settings.codex_cli_path  if provider.value == "codex_cli"  else
            None
        )

        # Persist an AIRun row up-front so it shows up in /ai-runs.
        run = models.AIRun(
            workflow=Workflow.EXTRACT_USER_PROFILE_FULL.value,
            provider=provider.value,
            status="queued",
        )
        db.add(run)
        db.commit()
        db.refresh(run)
        run_id = run.id

        base_ctx = _resolve_user_context(db)
        request = RunRequest(
            workflow=Workflow.EXTRACT_USER_PROFILE_FULL,
            params=base_ctx,
            timeout_s=420,
        )

        async def drain():
            async for _ in _run_and_stream(run_id, request, provider, cli_path):
                pass  # silently consume — _apply_workflow_result updates DB

        # If we're inside a running loop, schedule; otherwise create one.
        try:
            loop = asyncio.get_event_loop()
            if loop.is_running():
                asyncio.create_task(drain())
            else:
                loop.run_until_complete(drain())
        except RuntimeError:
            asyncio.run(drain())
    except Exception:
        pass
    finally:
        db.close()


@router.get("/{doc_id}")
def get_document(doc_id: int, db: Session = Depends(get_db)):
    d = db.get(models.Document, doc_id)
    if not d:
        raise HTTPException(404, "Document not found.")
    out = _doc_to_dict(d)
    out["text"] = d.text or ""
    return out


@router.get("/{doc_id}/file")
def download_document(doc_id: int, db: Session = Depends(get_db)):
    d = db.get(models.Document, doc_id)
    if not d or not d.file_path:
        raise HTTPException(404, "Document not found.")
    p = Path(d.file_path)
    if not p.exists():
        raise HTTPException(404, "File missing on disk.")
    return FileResponse(str(p), filename=p.name)


@router.patch("/{doc_id}")
def update_document(doc_id: int, payload: dict, db: Session = Depends(get_db)):
    d = db.get(models.Document, doc_id)
    if not d:
        raise HTTPException(404, "Document not found.")
    if "title" in payload and isinstance(payload["title"], str):
        d.title = payload["title"]
    if "is_default" in payload and bool(payload["is_default"]):
        db.query(models.Document).filter(
            models.Document.kind == d.kind,
            models.Document.id != d.id,
        ).update({"is_default": False})
        d.is_default = True
    elif "is_default" in payload and not payload["is_default"]:
        d.is_default = False
    d.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(d)
    return _doc_to_dict(d)


@router.delete("/{doc_id}", status_code=204)
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    d = db.get(models.Document, doc_id)
    if not d:
        raise HTTPException(404, "Document not found.")
    if d.file_path:
        p = Path(d.file_path)
        if p.exists():
            try:
                p.unlink()
            except OSError:
                pass
    db.delete(d)
    db.commit()
    return None
